import logging
import os
import re

from docx import Document
from docx.oxml import CT_P, CT_Tbl
from openpyxl import Workbook

from app.config.result_messages import ResultMessages
from app.constants.constants import account_number, currency_code
from app.constants.regex_patterns import ACCOUNT_NUMBER_PATTERN, CURRENCY_CODE_PATTERN, OPERATION_DATE_PATTERN, \
    PAYMENT_PURPOSE_PATTERN
from app.preprocessor.preprocessor import ProcessingError, ErrorSeverity


class AdditionalData:
    def __init__(self):
        self.account_number = ""
        self.currency_code = ""


logger = logging.getLogger(__name__)

def docx_to_xlsx(input_file, output_file) -> tuple[bool, list[ProcessingError]]:
    errors = []
    try:
        # Проверка существования входного файла
        if not os.path.exists(input_file):
            error = ProcessingError(
                code=ResultMessages.ERROR_DOCX_FILE_NOT_FOUND.status_code,
                message=ResultMessages.ERROR_DOCX_FILE_NOT_FOUND.message,
                severity=ErrorSeverity.CRITICAL,
                details={"file_path": input_file}
            )
            logger.error(error.message)
            errors.append(error)
            return False, errors


        # Читаем содержимое .docx через docx
        doc = Document(input_file)

        # Проверка, что документ не пустой
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            error = ProcessingError(
                code=ResultMessages.ERROR_DOCX_EMPTY_DOCUMENT.status_code,
                message=ResultMessages.ERROR_DOCX_EMPTY_DOCUMENT.message,
                severity=ErrorSeverity.WARNING,
                details={"file_path": input_file}
            )
            logger.warning(error.message)
            errors.append(error)

        # Создаем новый .xlsx файл
        wb = Workbook()
        ws = wb.active

        # Переменные для отслеживания строки в Excel
        row = 1
        # Массив с номерами счетов должников
        additional_datas = []
        # Проход по всем элементам в .docx
        for element in doc.element.body:
            # Проверяем, является ли элемент параграфом (CT_P)
            if isinstance(element, CT_P):
                paragraph = element.text.strip()
                if paragraph:
                    ws.cell(row=row, column=1, value=paragraph)
                    row += 1
            # Проверяем, является ли элемент таблицей (CT_Tbl)
            elif isinstance(element, CT_Tbl):
                table = next(t for t in doc.tables if t._tbl == element)
                for table_row in table.rows:
                    table_data = [cell.text.strip() for cell in table_row.cells]
                    # Ищем номер счета и код валюты
                    get_account_number_and_currency_code(additional_datas, table, table_data)
                    if len(table_data) < 4 and len(additional_datas) > 0:
                        intersect = [i.account_number for i in additional_datas if i.account_number in table_data]
                        if len(intersect) > 0:
                            if additional_datas[0].account_number not in intersect:
                                del additional_datas[0]
                    if len(table_data) > 5 and [i for i, s in enumerate(table_data) if
                                                re.search(OPERATION_DATE_PATTERN, s, re.IGNORECASE)] and [i for i, s in
                                                                                                          enumerate(table_data) if
                                                                                                          re.search(PAYMENT_PURPOSE_PATTERN, s,
                                                                                                                    re.IGNORECASE)]:
                        table_data.append(account_number)
                        table_data.append(currency_code)
                        for row_index in range(len(table.rows)):
                            ws.cell(row=row_index+row, column=len(table_data)-1, value=additional_datas[0].account_number)
                            ws.cell(row=row_index+row, column=len(table_data), value=additional_datas[0].currency_code)
                    # Проверяем, что есть данные в table_data
                    for col, cell_text in enumerate(table_data, start=1):
                        ws.cell(row=row, column=col, value=cell_text)
                    row += 1
                row += 1

        # Проверка возможности сохранения файла
        try:
            wb.save(output_file)
        except PermissionError:
            error = ProcessingError(
                code=ResultMessages.ERROR_DOCX_SAVE_PERMISSION_DENIED.status_code,
                message=ResultMessages.ERROR_DOCX_SAVE_PERMISSION_DENIED.message,
                severity=ErrorSeverity.CRITICAL,
                details={"output_file": output_file}
            )
            logger.error(error.message)
            errors.append(error)
            return False, errors
        except Exception as save_error:
            error = ProcessingError(
                code=ResultMessages.ERROR_DOCX_CONVERSION_FAILED.status_code,
                message=ResultMessages.ERROR_DOCX_CONVERSION_FAILED.message,
                severity=ErrorSeverity.CRITICAL,
                details={"exception": str(save_error)}
            )
            logger.error(error.message)
            errors.append(error)
            return False, errors

        logger.info(f"Файл {input_file} успешно конвертирован в формат .xlsx!")
        return True, errors

    except Exception as e:
        error = ProcessingError(
            code=ResultMessages.ERROR_DOCX_CONVERSION_FAILED.status_code,
            message=ResultMessages.ERROR_DOCX_CONVERSION_FAILED.message,
            severity=ErrorSeverity.CRITICAL,
            details={"exception": str(e), "file_path": input_file}
        )
        logger.error(error.message)
        errors.append(error)
        return False, errors


def get_account_number_and_currency_code(additional_datas, table, table_data):
    """
    Find and extract account number and currency code from table data.

    The function takes a list of table data and tries to find the index of
    the column which contains account number and currency code. It then
    iterates over each row in the table and extracts the account number and
    currency code if they are present.

    Args:
        additional_datas (list): A list of AdditionalData objects where
            the extracted data will be stored.
        table (docx.table.Table): The table from which the data is being
            extracted.
        table_data (list): A list of strings where each string is a cell
            in the table.

    Returns:
        None
    """
    if len(table_data) > 1:
        account_number_index = [i for i, s in enumerate(table_data) if re.search(ACCOUNT_NUMBER_PATTERN, s, re.IGNORECASE)]
        currency_code_index = [i for i, s in enumerate(table_data) if re.search(CURRENCY_CODE_PATTERN, s, re.IGNORECASE)]
        if account_number_index and currency_code_index:
            for current_table_row in table.rows:
                additional_data = AdditionalData()
                if re.search(r'\b\d{20}\b', current_table_row.cells[account_number_index[0]].text.strip(),
                             re.IGNORECASE):
                    additional_data.account_number = current_table_row.cells[account_number_index[0]].text.strip()
                if re.search(r'\b\d{3}\b', current_table_row.cells[currency_code_index[0]].text.strip(), re.IGNORECASE):
                    additional_data.currency_code = current_table_row.cells[currency_code_index[0]].text.strip()
                if additional_data.account_number or additional_data.currency_code:
                    additional_datas.append(additional_data)